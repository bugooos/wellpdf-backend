from fastapi import APIRouter, UploadFile, File, HTTPException
from fastapi.responses import FileResponse
import shutil
import os
import pandas as pd
from docx import Document

from core.utils import temp_input_path, safe_output_path

router = APIRouter()

@router.post("/convert/docx-to-xlsx", summary="Convert DOCX to XLSX")
async def docx_to_xlsx(file: UploadFile = File(...)):
    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(
            status_code=400,
            detail="Only DOCX files are supported (DOC is not supported)"
        )

    input_path = temp_input_path("docx")
    output_path = safe_output_path(file.filename, "xlsx")

    with open(input_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    try:
        doc = Document(input_path)

        writer = pd.ExcelWriter(output_path, engine="openpyxl")

        if doc.tables:
            # ✅ Multiple sheets (one per table)
            for idx, table in enumerate(doc.tables):
                rows = []
                for row in table.rows:
                    rows.append([cell.text.strip() for cell in row.cells])

                df = pd.DataFrame(rows)
                df.to_excel(
                    writer,
                    sheet_name=f"Table_{idx + 1}",
                    index=False,
                    header=False
                )
        else:
            # ✅ Fallback: paragraphs → single sheet
            rows = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    rows.append([text])

            if not rows:
                raise HTTPException(
                    status_code=400,
                    detail="No readable content found in DOCX file"
                )

            df = pd.DataFrame(rows)
            df.to_excel(
                writer,
                sheet_name="Content",
                index=False,
                header=False
            )

        writer.close()

        return FileResponse(
            output_path,
            filename=os.path.basename(output_path),
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

    except HTTPException:
        raise

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    finally:
        if os.path.exists(input_path):
            os.remove(input_path)
