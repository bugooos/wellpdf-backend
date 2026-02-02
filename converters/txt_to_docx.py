from fastapi import APIRouter, UploadFile, File, HTTPException
from fastapi.responses import FileResponse
from docx import Document
import shutil, os, re

from core.utils import temp_input_path, safe_output_path

router = APIRouter()

# Remove XML-invalid characters
INVALID_XML_CHARS = re.compile(
    r"[\x00-\x08\x0B\x0C\x0E-\x1F]"
)

def clean_text(text: str) -> str:
    return INVALID_XML_CHARS.sub("", text)


@router.post("/convert/txt-to-docx", summary="Convert TXT to DOCX")
async def txt_to_docx(file: UploadFile = File(...)):
    if not file.filename.lower().endswith(".txt"):
        raise HTTPException(status_code=400, detail="Only TXT files allowed")

    input_path = temp_input_path("txt")
    output_path = safe_output_path(file.filename, "docx")

    with open(input_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    try:
        doc = Document()

        with open(input_path, "r", encoding="utf-8", errors="ignore") as f:
            for raw_line in f:
                line = clean_text(raw_line.rstrip())
                if line.strip():
                    doc.add_paragraph(line)
                else:
                    doc.add_paragraph("")  # preserve blank lines

        doc.save(output_path)

        return FileResponse(
            output_path,
            filename=os.path.basename(output_path),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    finally:
        if os.path.exists(input_path):
            os.remove(input_path)
